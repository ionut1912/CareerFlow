from __future__ import annotations

import hashlib
import shutil
import subprocess
import tempfile
import threading
import time
from dataclasses import dataclass
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}

_content_cache: dict[str, tuple[DocumentContent, float]] = {}
_cache_lock = threading.Lock()
_CACHE_TTL = 600


@dataclass(frozen=True)
class DocumentContent:
    filename: str
    total_pages: int
    text: str
    pages: list[str]


def file_hash(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def _evict_stale() -> None:
    now = time.monotonic()
    stale = [k for k, (_, ts) in _content_cache.items() if now - ts > _CACHE_TTL]
    for k in stale:
        _content_cache.pop(k, None)


def get_cached_content(key: str) -> DocumentContent | None:
    with _cache_lock:
        _evict_stale()
        entry = _content_cache.get(key)
        if entry:
            return entry[0]
    return None


def set_cached_content(key: str, content: DocumentContent) -> None:
    with _cache_lock:
        _evict_stale()
        _content_cache[key] = (content, time.monotonic())


def extract_text_from_pdf(path: Path) -> DocumentContent:
    pages: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    return DocumentContent(
        filename=path.name,
        total_pages=len(pages),
        text="\n\n".join(pages),
        pages=pages,
    )


def extract_text_from_docx(path: Path) -> DocumentContent:
    doc = DocxDocument(str(path))
    raw_paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                raw_paragraphs.append(row_text)
    chunk_size = 30
    pages: list[str] = [
        "\n".join(raw_paragraphs[i : i + chunk_size])
        for i in range(0, max(len(raw_paragraphs), 1), chunk_size)
    ] or [""]
    return DocumentContent(
        filename=path.name,
        total_pages=len(pages),
        text="\n\n".join(pages),
        pages=pages,
    )


def convert_doc_to_docx(path: Path) -> Path:
    out_dir = Path(tempfile.mkdtemp())
    subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", str(out_dir), str(path)],
        check=True,
        capture_output=True,
        timeout=30,
    )
    converted = out_dir / (path.stem + ".docx")
    if not converted.exists():
        raise RuntimeError(f"LibreOffice failed to convert {path.name}")
    return converted


def extract_text_from_document(path: str | Path) -> DocumentContent:
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    key = file_hash(path)
    cached = get_cached_content(key)
    if cached is not None:
        return cached

    ext = path.suffix.lower()
    if ext == ".pdf":
        content = extract_text_from_pdf(path)
    elif ext == ".docx":
        content = extract_text_from_docx(path)
    elif ext == ".doc":
        converted = convert_doc_to_docx(path)
        try:
            content = extract_text_from_docx(converted)
        finally:
            shutil.rmtree(converted.parent, ignore_errors=True)
    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    set_cached_content(key, content)
    return content


def chunk_text(text: str, max_chars: int = 12000) -> list[str]:
    if len(text) <= max_chars:
        return [text]
    chunks: list[str] = []
    paragraphs = text.split("\n\n")
    current: list[str] = []
    current_len = 0
    for para in paragraphs:
        if current_len + len(para) + 2 > max_chars and current:
            chunks.append("\n\n".join(current))
            current = []
            current_len = 0
        current.append(para)
        current_len += len(para) + 2
    if current:
        chunks.append("\n\n".join(current))
    return chunks